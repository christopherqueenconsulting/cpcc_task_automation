import os
import os.path
import tempfile
import time
import zipfile
from enum import Enum, StrEnum
from functools import lru_cache
from random import randint
from typing import Optional, Annotated, List, Union

import docx
import mammoth
# import markdownify
import pandas as pd
import textract
from bs4 import BeautifulSoup
from cqc_cpcc.utilities.date import get_datetime
from cqc_cpcc.utilities.env_constants import IS_GITHUB_ACTION
from cqc_cpcc.utilities.logger import logger
from cqc_cpcc.utilities.selenium_util import (
    get_driver_wait,
    click_element_wait_retry,
    take_and_show_screenshot,
    wait_for_user_action,
)
from docx import Document
from markdownify import markdownify as md
from ordered_set import OrderedSet
from pydantic import BaseModel, Field, StrictStr, PositiveInt
from selenium.common import (
    TimeoutException,
    StaleElementReferenceException,
    ElementNotInteractableException,
    NoSuchElementException,
)
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

# from simplify_docx import simplify

# Global Constants
LINE_DASH_COUNT = 33


class Satisfactory(Enum):
    YES = 1
    NO = 0


def are_you_satisfied():
    """Prompts the user to select if they are satisfied or not."""

    enum = Satisfactory

    logger.info("Are you satisfied?")
    for i, member in enumerate(enum):
        logger.info("%s: %s", member.value, member.name)

    default = Satisfactory.YES
    default_value = default.value
    user_input = int(input('Enter your selection [' + str(default_value) + ']: ').strip() or default_value)

    try:
        sf = Satisfactory(user_input)
        logger.info("You selected %s", sf.name)
        return sf.value == default_value
    except ValueError:
        logger.warning("Invalid selection.")
        return are_you_satisfied() == default_value


def first_two_uppercase(string):
    """Returns the first two letters of a string, in uppercase."""
    return string[:2].upper()


def html_table_to_dict(table_text) -> [dict, dict]:
    # logger.info('Table: %s' %table_text)
    table_text = table_text.replace("\n", " ")  # Remove new lines
    # TODO: ??? Add Pipe seperator to <spans> or <p>
    table_text = table_text.replace("</p>", "|</p>")  # This is to be able to split assignment names later
    soup = BeautifulSoup(table_text, 'html.parser')

    rows = []
    headers = []

    for tr in soup.find('table').find('tbody').find_all('tr'):
        row = []

        for th in tr.find_all('th'):
            try:
                headers.append(th.text.strip())
            except:
                continue

        for td in tr.find_all('td'):
            try:
                row.append(td.text.strip())
            except:
                continue
        # logger.info('Row Length : %s' % str(len(row)))
        if 1 < len(row) < 3:
            # Modify the 2nd and 3rd column so its has the same date text
            row_date = get_datetime(row[1].split("-")[0].strip())
            row[1] = row_date.strftime("%Y-%m-%d")
            row.append(row[1])
        if len(row) == 3:
            rows.append(row)

    # logger.info('Table Headers: %s' % headers)
    # logger.info('Table (extracted): %s' % rows)

    df = pd.DataFrame(rows, columns=headers)

    return headers, df.to_dict()


def get_unique_names(list_names: list):
    # insert the list to the set
    list_set = set(list_names)
    # convert the set to the list
    unique_list = (list(list_set))
    # Sort alphabetically
    unique_list.sort()

    return unique_list


def get_unique_names_flip_first_last(list_names: list) -> list:
    unique_names = get_unique_names(list_names)
    names_flipped = map(lambda kv: flip_name(kv), unique_names)
    return list(names_flipped)


def flip_name(full_name: str):
    separator = ','
    name_parts = full_name.split(separator)
    name_parts.reverse()
    return separator.join(name_parts)


"""
def get_html_as_markdown(html: str, code_language='java') -> str:
    # convert html to markdown

    markdown = markdownify.markdownify(html, heading_style="ATX", code_language=code_language)
    return markdown
"""


class ExtendedEnum(StrEnum):

    @classmethod
    def list(cls) -> list:
        return [c.value for c in cls]


class CodeError(BaseModel):
    """Object representing a Code Error"""
    error_type: Optional[
        Annotated[
            Enum,
            Field(description="The type of coding error")
        ]
    ] = None
    code_error_lines: Optional[
        Annotated[
            List[StrictStr],
            Field(description="An array list of the lines of code that are relevant to the coding error")
        ]
    ] = None

    line_numbers_of_error_holder: Optional[
        Annotated[
            List[PositiveInt],
            Field(description="The list of line numbers relevant to this coding error")
        ]
    ] = None

    error_details: Annotated[
        StrictStr,
        Field(description="The details about this coding error")
    ]

    def set_line_numbers_of_error(self, line_numbers_of_errors: List[PositiveInt]):
        self.line_numbers_of_error_holder = line_numbers_of_errors

    @property
    def line_numbers_of_errors(self) -> List:
        if self.line_numbers_of_error_holder is None:
            return []
        else:
            return sorted(set(self.line_numbers_of_error_holder))

    def __str__(self):
        import cqc_cpcc.utilities.env_constants as EC
        lines_string_complete = ""
        if EC.SHOW_ERROR_LINE_NUMBERS and self.line_numbers_of_errors is not None:
            lines_string = ", ".join(map(str, self.line_numbers_of_errors))
            lines_string_complete = f"\n\tOn Line(s) #: {lines_string}"
        error_details_string = "\t" + self.error_details.replace("\n", "\n\t")
        return f"{self.error_type.value}:{lines_string_complete}\n{error_details_string}"


class ErrorHolder(BaseModel):

    def get_combined_errors_by_type(self, code_errors: List[CodeError]) -> List[CodeError]:
        errors = {}
        final_errors = []
        for code_error in code_errors:
            if code_error.error_type not in errors:
                errors[code_error.error_type] = []
            errors[code_error.error_type].append(code_error)

        # print("Errors by Type")
        # pprint(errors)

        for error_type in errors:
            code_errors_by_type = errors[error_type]
            nested_list = [x.line_numbers_of_errors for x in code_errors_by_type]
            # print("Nested List")
            # pprint(nested_list)
            # Use map and lambda to flatten the nested list
            # Use list comprehension to flatten the nested list
            try:
                flattened_list = [number for sublist in nested_list for number in sublist]
            except TypeError as e:
                flattened_list = []

            # print("Flattened List")
            # pprint(flattened_list)
            unique_sorted_list = sorted(set(flattened_list))
            # print("Unique Sorted List")
            # pprint(unique_sorted_list)
            # error_details_list = list(map(lambda x: x.error_details, code_errors_by_type))

            # Combine the error details
            error_details_list = [x.error_details for x in code_errors_by_type]
            unique_error_details_list = list(OrderedSet(error_details_list))
            error_details = "\n".join(unique_error_details_list)

            # Combine the code_error_lines
            try:
                code_error_lines_list = [x.code_error_lines for x in code_errors_by_type]
                code_error_lines_flattened_list = [line for code_error_lines_sublist in code_error_lines_list for line
                                                   in
                                                   code_error_lines_sublist]
            except TypeError as e:
                code_error_lines_flattened_list = []

            final_errors.append(
                CodeError(error_type=error_type,
                          line_numbers_of_error=unique_sorted_list,
                          # TODO: Do we need this since the line numbers come after from code???
                          error_details=error_details,
                          code_error_lines=code_error_lines_flattened_list))

        # print("Final Errors by Type Combined")
        # pprint(final_errors)

        return final_errors


def wrap_code_in_markdown_backticks(code: str, code_type: str = "java") -> str:
    backticks = "`"
    code_fence = backticks * 3
    # See if code fence exist in code already
    while code_fence in code:
        code_fence += backticks

    prefix = code_fence + code_type + "\n"
    suffix = "\n" + code_fence
    return prefix + code.strip() + suffix


def merge_lists(list1, list2):
    """Merges two lists, where one might be None.

    Args:
      list1: The first list.
      list2: The second list.

    Returns:
      A merged list.
    """
    if list1 is None and list2 is None:
        return None
    elif list1 is None:
        return list2
    elif list2 is None:
        return list1
    else:
        return list1 + list2


def convert_tables_to_json_in_tmp__file(doc: Document) -> str:
    for table in doc.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data)

        # Remove the table
        t = table._element
        parent = t.getparent()
        parent.remove(t)

        # Add new json string to the parent in its place
        doc.add_paragraph(df.to_json(orient="records"))

        # Clear the table reference
        t._t = t._element = None

    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)

    return temp_file.name


@lru_cache(maxsize=None)
def convert_content_to_markdown(content: str) -> str:
    return md(content)


@lru_cache(maxsize=None)
def convert_xlsx_to_markdown(file_path: str) -> str:
    """Convert Excel sheets into well-formatted markdown."""
    try:
        sheets = pd.read_excel(file_path, sheet_name=None)
        markdown_output = []

        for sheet_name, df in sheets.items():
            # Process each column individually based on its dtype
            for column in df.columns:
                if pd.api.types.is_numeric_dtype(df[column]):
                    # Replace numeric NaNs with a placeholder or keep as is
                    df[column] = df[column].apply(lambda x: '' if pd.isna(x) else x)
                else:
                    df[column] = df[column].fillna('')

            markdown = df.to_markdown(index=False, tablefmt='github')
            markdown_output.append(f"### {sheet_name}\n\n{markdown}")

        return "\n\n".join(markdown_output)
    except Exception as e:
        return f"Error converting Excel file to markdown: {str(e)}"


@lru_cache(maxsize=None)
def read_file(file_path: str, convert_to_markdown: bool = False) -> str:
    """ Return the file contents in string format.
    
    For PDF files (.pdf): Extracts text using PyMuPDF/PyPDF to avoid binary data
    For audio files (.mp3, .wav, .m4a, .ogg): Transcribes using OpenAI Whisper
    For video files (.mp4, .avi, .mov, .webm): Returns metadata and grading instructions
    For HTML files: Extracts text content (removes scripts/styles)
    For other files: Returns text content as-is
    """
    file_name, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    # If file is PDF, extract text using specialized PDF library
    if file_extension == '.pdf':
        from cqc_cpcc.utilities.pdf_utils import extract_text_from_pdf
        contents = extract_text_from_pdf(file_path)
    elif convert_to_markdown:
        with open(file_path, mode='rb') as f:
            # results = mammoth.convert_to_markdown(f)
            results = mammoth.convert_to_html(f)
            contents = convert_content_to_markdown(results.value)
        # contents = results.value
    # If file is HTML, extract text content
    elif file_extension in ['.html', '.htm']:
        with open(file_path, mode='r', encoding='utf-8') as f:
            html_content = f.read()
        # Parse HTML and extract text
        soup = BeautifulSoup(html_content, 'html.parser')
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        # Get text content
        text = soup.get_text()
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        contents = '\n'.join(chunk for chunk in chunks if chunk)
    # If file is audio, transcribe it using OpenAI Whisper
    elif file_extension in ['.mp3', '.wav', '.m4a', '.ogg']:
        try:
            # Import asyncio to run the async transcription
            import asyncio
            from cqc_cpcc.utilities.AI.openai_client import transcribe_audio, format_transcription_for_grading

            # Run the async transcription - handle event loop properly
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = None

            if loop is not None and loop.is_running():
                # We're in an async context - create new loop in thread
                import concurrent.futures
                def run_in_thread():
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(transcribe_audio(file_path))
                    finally:
                        new_loop.close()

                with concurrent.futures.ThreadPoolExecutor() as pool:
                    future = pool.submit(run_in_thread)
                    transcription = future.result()
            else:
                # No running loop - create and use one
                transcription = asyncio.run(transcribe_audio(file_path))

            contents = format_transcription_for_grading(transcription)
        except Exception as e:
            # If transcription fails, return error message with file info
            file_size = os.path.getsize(file_path) / (1024 * 1024)
            contents = f"""[AUDIO FILE: {os.path.basename(file_path)}]
File type: {file_extension[1:].upper()}
File size: {file_size:.2f} MB

Error: Failed to transcribe audio file: {str(e)}
Please manually review this audio file for grading."""
    # If file is video, return metadata and instructions
    elif file_extension in ['.mp4', '.avi', '.mov', '.webm']:
        try:
            import asyncio
            from cqc_cpcc.utilities.AI.openai_client import process_video_file

            # Run the async video processing - handle event loop properly
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = None

            if loop is not None and loop.is_running():
                # We're in an async context - create new loop in thread
                import concurrent.futures
                def run_in_thread():
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(process_video_file(file_path))
                    finally:
                        new_loop.close()

                with concurrent.futures.ThreadPoolExecutor() as pool:
                    future = pool.submit(run_in_thread)
                    contents = future.result()
            else:
                # No running loop - create and use one
                contents = asyncio.run(process_video_file(file_path))
        except Exception as e:
            file_size = os.path.getsize(file_path) / (1024 * 1024)
            contents = f"""[VIDEO FILE: {os.path.basename(file_path)}]
File type: {file_extension[1:].upper()}
File size: {file_size:.2f} MB

Error: Failed to process video file: {str(e)}
Please manually review this video file for grading."""
    # If file ends in .xlsx, convert it to markdown
    elif file_extension in ['.xlsx', '.xls', '.xlsm']:
        contents = convert_xlsx_to_markdown(file_path)
    # If file ends in .docx will convert it to json and return
    elif file_extension == ".docx":
        # read in a document
        my_doc = docx.Document(file_path)

        # Find any tables and replace with json strings
        tmp_file = convert_tables_to_json_in_tmp__file(my_doc)

        # coerce to JSON using the standard options

        # contents = simplify(my_doc)

        # contents = textract.parsers.process(file_path)
        # print("Extracting contents from: %s" % tmp_file)
        contents = textract.process(tmp_file).decode('utf-8')
        os.remove(tmp_file)

    else:

        encodings = ['utf-8', 'latin-1', 'utf-16', 'ascii']
        contents = ""
        index = 0

        while index < len(encodings):
            try:
                with open(file_path, mode='r', encoding=encodings[index]) as f:
                    contents = f.read()
                break
            except UnicodeDecodeError:
                index += 1
            except Exception:
                contents = ""
                break

    return str(contents)


def read_files(file_paths: Union[str, List[str]], convert_to_markdown: bool = False) -> str:
    if isinstance(file_paths, str):
        # If a single string is provided, treat it as a file path
        return read_file(file_paths, convert_to_markdown)
    elif isinstance(file_paths, list):
        # If a list is provided, loop through each file path and read the file
        concatenated_content = ""
        for path in file_paths:
            file_content = read_file(path, convert_to_markdown)
            concatenated_content += file_content + "\n\n"  # You can customize the separator if needed
        return concatenated_content
    else:
        return "Invalid input. Please provide a string or a list of strings (file paths)."


def dict_to_markdown_table(data, headers):
    # Create the header row
    markdown_table = "| " + " | ".join(headers) + " |\n"
    markdown_table += "| " + " | ".join(["-" * len(header) for header in headers]) + " |\n"

    # Iterate over the dictionary items and add rows to the table
    for row_data in data:
        markdown_table += "| " + " | ".join([str(row_data.get(header, '')) for header in headers]) + " |\n"

    return markdown_table


def extract_and_read_zip(file_path: str, accepted_file_types: list[str]) -> dict:
    unacceptable_file_prefixes = ['._']
    # BrightSpace export files that should never be graded
    unacceptable_file_names = {'index.html', 'index.htm'}
    students_data = {}

    if file_path.endswith('.zip'):

        # Open the zip file
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            # Iterate over each file in the zip archive
            for file_info in zip_ref.infolist():
                # Extract the file name and directory name
                file_name = os.path.basename(file_info.filename)
                directory_name = os.path.dirname(file_info.filename)

                # Check if the directory name represents a student folder
                folder_name_delimiter = ' - '
                if directory_name and folder_name_delimiter in directory_name:
                    student_name = directory_name.split(folder_name_delimiter)[1]

                    # Check if the file has an accepted file type
                    # Skip BrightSpace export files (case-insensitive)
                    if (file_name.endswith(tuple(accepted_file_types))
                            and not file_name.startswith(tuple(unacceptable_file_prefixes))
                            and file_name.lower() not in unacceptable_file_names):
                        # Read the file contents
                        with zip_ref.open(file_info.filename) as file:
                            # TODO: Change to modules on read file method
                            # file_contents = file.read().decode('utf-8')  # Assuming UTF-8 encoding
                            sub_file_name, sub_file_extension = os.path.splitext(file_name)
                            prefix = 'from_zip_' + str(randint(1000, 100000000)) + "_"
                            temp_file = tempfile.NamedTemporaryFile(delete=False, prefix=prefix,
                                                                    suffix=sub_file_extension)
                            temp_file.write(file.read())
                            # file_contents = read_file(
                            #    temp_file.name)  # Reading this way incase it may be .docx or some other type we want to pre-process differently

                        # Store the file contents in the dictionary
                        if student_name not in students_data:
                            students_data[student_name] = {}
                        # students_data[student_name][file_name] = file_contents
                        students_data[student_name][file_name] = temp_file.name

    return students_data


# Seconds to let the MFA prompt finish its slide-in animation before we capture
# the screenshot / matching number, so the number isn't caught mid-transition.
MFA_SCREENSHOT_SETTLE_SECONDS = 2.5

# Max seconds to wait for the number-matching challenge to render after Sign in
# before concluding MFA wasn't required (e.g. a cached session skipped it).
MFA_DETECT_TIMEOUT_SECONDS = 8

# Max seconds to wait for the user to approve the MFA prompt on their device
# before giving up. The "Stay signed in?" page only appears after approval, so
# this must comfortably exceed the time to notice + open the app + approve.
MFA_APPROVAL_TIMEOUT_SECONDS = 120


class MfaCancelled(Exception):
    """Raised when the user cancels a pending MFA approval (web app)."""


def _mfa_settle_before_capture() -> None:
    """Pause briefly so the MFA prompt animation settles before capture."""
    time.sleep(MFA_SCREENSHOT_SETTLE_SECONDS)


def _publish_mfa_challenge(driver: WebDriver, context: str, mfa_handler, message: str) -> str | None:
    """Capture the current matching number + screenshot and publish to the handler.

    Returns the captured number (or ``None``). Best-effort: never raises so it
    can't break login or the approval-wait loop.
    """
    try:
        from cqc_cpcc.utilities.selenium_util import capture_mfa_challenge
        challenge = capture_mfa_challenge(driver, context)
        challenge.message = message
        mfa_handler.on_challenge(challenge)
        return challenge.number
    except Exception as e:  # noqa: BLE001 - never let MFA notify break login
        logger.warning("Could not notify MFA handler: %s", e)
        return None


def _notify_mfa(driver: WebDriver, context: str, mfa_handler, message: str) -> None:
    """Surface an MFA number-matching prompt.

    When ``mfa_handler`` is provided (e.g. the headless web app's ``MfaBridge``),
    capture the on-screen matching number + screenshot and hand them off so the
    UI can display them on the same page. The number is re-published continuously
    while we wait for approval (see ``_wait_for_mfa_approval``), so an initial
    miss (number still animating in) is corrected on the next poll. Otherwise
    fall back to the existing CLI behavior: open a screenshot on the user's real
    display and log the instruction — and also extract + print the number.
    """
    # The matching number slides into the center; wait for it to settle so the
    # screenshot/number aren't captured mid-animation.
    _mfa_settle_before_capture()

    if mfa_handler is not None:
        number = _publish_mfa_challenge(driver, context, mfa_handler, message)
        # Never log the raw matching number — it's an auth challenge value. It's
        # surfaced to the user via the MFA handler + on-screen screenshot, not the
        # server log. Log only whether it has been read yet.
        logger.info(
            "🔐 MFA number matching%s — %s",
            "" if number else " (number not read yet; will retry)",
            message,
        )
    else:
        from cqc_cpcc.utilities.selenium_util import describe_mfa_dom, extract_mfa_number
        number = extract_mfa_number(driver)
        take_and_show_screenshot(driver, f"{context}_mfa")
        if number:
            # Displayed on the screenshot shown to the user; not logged.
            logger.info("🔐 MFA matching number read from page (shown on screen).")
        else:
            # Selector didn't match — dump candidates so the selector can be tuned.
            logger.info("Could not read the MFA number from the page selectors.")
            logger.info(describe_mfa_dom(driver))
        logger.info(message)


def _wait_for_mfa_approval(
        driver: WebDriver,
        display_xpath: str,
        mfa_handler=None,
        timeout: int = MFA_APPROVAL_TIMEOUT_SECONDS,
        context: str = "",
        message: str = "",
) -> None:
    """Block until the MFA number-matching screen clears (user approved).

    Polls for the prompt element to disappear, which signals approval, so the
    login can proceed to the post-MFA page. While the prompt is still present and
    an ``mfa_handler`` is provided, the matching number + screenshot are
    re-published every poll so the web-app page reliably shows the current number
    (even if it wasn't readable on the first capture). Honors
    ``mfa_handler.cancelled`` so the web app's Cancel button can abort. Returns
    quietly on timeout (the caller then attempts the next step, which will
    surface a clearer error if needed).
    """
    deadline = time.time() + timeout
    last_number = None
    while time.time() < deadline:
        if mfa_handler is not None and getattr(mfa_handler, "cancelled", False):
            raise MfaCancelled("MFA approval cancelled by user")
        try:
            still_present = bool(driver.find_elements(By.XPATH, display_xpath))
        except Exception:  # noqa: BLE001
            still_present = False
        if not still_present:
            logger.info("MFA prompt cleared — approval detected.")
            return
        # Re-publish so the UI keeps an up-to-date number/screenshot. This fixes
        # the case where the number was still animating in at first capture.
        if mfa_handler is not None:
            number = _publish_mfa_challenge(driver, context, mfa_handler, message)
            if number and number != last_number:
                last_number = number
                # Republished to the UI/handler above; the value itself is not
                # logged (auth challenge value — shown on screen instead).
                logger.info("🔐 MFA matching number updated — see the on-screen prompt.")
        time.sleep(1)
    logger.warning("Timed out waiting for MFA approval after %ss.", timeout)


# Max seconds to wait for the "Stay signed in?" (KMSI) prompt to appear before
# concluding the tenant skipped it. Kept short so login isn't delayed when absent.
KMSI_PROMPT_TIMEOUT_SECONDS = 3


def _dismiss_stay_signed_in(driver: WebDriver, wait_short, wait_long) -> None:
    """Click "No" on the "Stay signed in?" page if it appears, else skip fast.

    The prompt is optional (some tenants skip it / we may already be redirected
    back to the app after MFA), so this probes briefly with a direct presence
    check rather than a long element wait that would hang when it never shows.
    """
    kmsi_no_xpath = "//input[contains(@class, 'button-secondary') and contains(@value,'No')]"
    deadline = time.time() + KMSI_PROMPT_TIMEOUT_SECONDS
    while time.time() < deadline:
        if driver.find_elements(By.XPATH, kmsi_no_xpath):
            try:
                no_btn = click_element_wait_retry(
                    driver, wait_short, kmsi_no_xpath,
                    "Clicking 'No' to Stay Signed in", max_try=0)
                wait_long.until(
                    EC.invisibility_of_element(no_btn),
                    'Waiting for login to be successful')
            except (TimeoutException, StaleElementReferenceException,
                    ElementNotInteractableException, NoSuchElementException):
                logger.info("'Stay signed in?' prompt vanished before dismissal.")
            return
        time.sleep(0.5)
    logger.info("No 'Stay signed in?' prompt — login already completed.")


def _resolve_mfa(mfa_handler) -> None:
    """Signal the MFA handler that login completed (UI can close the prompt)."""
    if mfa_handler is not None:
        try:
            mfa_handler.on_resolved()
        except Exception as e:  # noqa: BLE001
            logger.debug("MFA handler on_resolved failed: %s", e)


def login_if_needed(driver: WebDriver, mfa_handler=None):
    # sleep for 3 seconds
    time.sleep(3)
    if "Web Login Service" in driver.title:
        # Duo Login
        duo_login(driver, mfa_handler=mfa_handler)
    elif "Sign in to your account" in driver.title:
        # Microsoft Login
        microsoft_login(driver, mfa_handler=mfa_handler)


def _get_microsoft_account_tile_xpath(instructor_user_id: str, instructor_email: str) -> str:
    """Build an XPath that matches the instructor account tile on Microsoft account picker."""
    return (
            '//*[@id="tilesHolder"]//*[contains(text(), "'
            + instructor_user_id.lower()
            + '") or contains(text(), "'
            + instructor_email.lower()
            + '")]'
    )


def _get_microsoft_expected_display_name_xpath(instructor_user_id: str, instructor_email: str) -> str:
    """Build an XPath that matches an already-selected expected Microsoft account."""
    return (
            "//div[@id='displayName' and (contains(@title,'"
            + instructor_user_id.lower()
            + "') or contains(@title,'"
            + instructor_email.lower()
            + "'))]"
    )


def _is_xpath_present(driver: WebDriver, xpath: str) -> bool:
    """Fast presence probe that avoids waiting when a branch is clearly unavailable."""
    return bool(driver.find_elements(By.XPATH, xpath))


def _is_xpath_present_with_short_wait(wait_probe: WebDriverWait, xpath: str, message: str) -> bool:
    """Short wait probe used when an element may appear after a quick transition."""
    try:
        wait_probe.until(
            EC.presence_of_element_located((By.XPATH, xpath)),
            message,
        )
        return True
    except TimeoutException:
        return False


def _wait_for_microsoft_identifier_step(wait_short: WebDriverWait) -> None:
    """Wait until either account display or username input is present."""
    wait_short.until(
        lambda d: d.find_elements(By.NAME, "loginfmt") or d.find_elements(By.ID, "displayName"),
        "Waiting for Microsoft identifier step",
    )


def _click_use_another_account_if_present(
        driver: WebDriver,
        wait: WebDriverWait,
        wait_short: WebDriverWait,
        wait_probe: WebDriverWait,
) -> bool:
    """Click "Use another account" when available and wait for the identifier step."""
    use_another_xpath = "//div[contains(text(),'Use another')]/parent::div"

    if not _is_xpath_present(driver, use_another_xpath) and not _is_xpath_present_with_short_wait(
            wait_probe,
            use_another_xpath,
            "Waiting briefly for use another account button",
    ):
        return False

    click_element_wait_retry(
        driver,
        wait_probe,
        use_another_xpath,
        "Waiting for use another account button",
    )

    _wait_for_microsoft_identifier_step(wait_short)
    return True


def _resolve_microsoft_account_path(
        driver: WebDriver,
        wait: WebDriverWait,
        wait_short: WebDriverWait,
        wait_probe: WebDriverWait,
        instructor_user_id: str,
        instructor_email: str,
) -> bool:
    """Resolve Microsoft account selection path and return True when username entry can be skipped."""
    pick_account_xpath = _get_microsoft_account_tile_xpath(instructor_user_id, instructor_email)
    expected_display_name_xpath = _get_microsoft_expected_display_name_xpath(instructor_user_id, instructor_email)
    any_display_name_xpath = "//div[@id='displayName']"

    if _is_xpath_present(driver, expected_display_name_xpath):
        return True
    if _is_xpath_present(driver, pick_account_xpath) or _is_xpath_present_with_short_wait(
            wait_probe,
            pick_account_xpath,
            "Waiting briefly to see if pick account is visible",
    ):
        click_element_wait_retry(driver, wait_probe, pick_account_xpath, "Waiting for pick account selection")
        return True
    if _is_xpath_present(driver, any_display_name_xpath) or _is_xpath_present_with_short_wait(
            wait_probe,
            any_display_name_xpath,
            "Waiting briefly to see if another username is present",
    ):
        click_element_wait_retry(driver, wait_probe, "//button[@class='backButton']", "Waiting for back button")
        _click_use_another_account_if_present(driver, wait, wait_short, wait_probe)
        return False

    _click_use_another_account_if_present(driver, wait, wait_short, wait_probe)

    return False


def microsoft_login(driver: WebDriver, mfa_handler=None):
    # Enter in user info and password
    instructor_user_id = os.environ["INSTRUCTOR_USERID"]
    instructor_email = os.getenv("INSTRUCTOR_EMAIL", f"{instructor_user_id}@cpcc.edu")
    instructor_password = os.environ["INSTRUCTOR_PASS"]

    wait = get_driver_wait(driver, 15)
    wait_short = get_driver_wait(driver, 5)
    wait_probe = get_driver_wait(driver, 1)
    wait_long = get_driver_wait(driver, 30)

    original_window = driver.current_window_handle

    logger.info("Microsoft login page detected — checking for existing account selection or username entry.")

    # Wait for the title to change
    wait_short.until(EC.title_is("Sign in to your account"))

    logger.info("Microsoft login page loaded — proceeding with authentication.")

    wait_short.until(
        EC.presence_of_element_located((By.XPATH, "//div[@id='loginHeader']")),
        "Waiting for login screen presence",
    )

    logger.info("Checking for existing account selection or username entry.")

    # Username may already be prefilled; if a different user is shown, clear it.
    user_name_already_entered = _resolve_microsoft_account_path(
        driver,
        wait,
        wait_short,
        wait_probe,
        instructor_user_id,
        instructor_email,
    )

    if not user_name_already_entered:
        logger.info("Username Not already entered")

        # Enter username / email
        # username_field = driver.find_element(By.NAME, "loginfmt")
        username_field = wait_short.until(EC.element_to_be_clickable((By.NAME, "loginfmt")), "Waiting for username input")

        # Note : Avoid plain send_keys
        username_field.click()
        username_field.clear()
        username_field.send_keys(instructor_user_id + "@cpcc.edu")


        # Click Next
        click_element_wait_retry(driver, wait,
                                 "//input[contains(@class, 'button_primary') and contains(@value,'Next')]",
                                 "Waiting for Next Button", By.XPATH)

    logger.info("Wating for password input")

    # Enter password
    # password_field = driver.find_element(By.NAME, "passwd")
    password_field = wait_short.until(EC.element_to_be_clickable((By.NAME, "passwd")), "Waiting for password input")
    # Note : Avoid plain send_keys
    password_field.click()
    password_field.clear()
    password_field.send_keys(instructor_password)

    logger.info("Waiting for Sign in Button")

    # Click Sign In
    click_element_wait_retry(driver, wait, "//input[contains(@class, 'button_primary') and contains(@value,'Sign in')]",
                             "Waiting for Sign in Button", By.XPATH)

    # A Microsoft Authenticator "number matching" challenge may appear here. When
    # running headless (web app) the instructor can't see the browser, so surface
    # the matching number + screenshot via the MFA handler so they can approve it
    # on their phone. Detection is best-effort; absence is fine (push/no-MFA).
    number_match_xpath = "//*[@id='idRichContext_DisplaySign']"
    # The number-matching screen can take a few seconds to render after Sign in,
    # so use a longer wait than the 1s account-picker probe. presence_of_element
    # returns as soon as it appears, so the full timeout is only spent when MFA is
    # genuinely absent (e.g. a cached session that skipped the challenge).
    wait_mfa = get_driver_wait(driver, MFA_DETECT_TIMEOUT_SECONDS)
    if _is_xpath_present_with_short_wait(
            wait_mfa,
            number_match_xpath,
            "Checking for Microsoft number-matching challenge",
    ):
        mfa_message = "Enter the number shown into your Microsoft Authenticator app, then approve."
        _notify_mfa(driver, "microsoft", mfa_handler, mfa_message)
        # The "Stay signed in?" page only appears after the user approves, which
        # can take a while — wait for the number prompt to clear before moving on.
        # Pass context/message so the number + screenshot are re-published to the
        # web-app page every poll until approval.
        _wait_for_mfa_approval(
            driver, number_match_xpath, mfa_handler,
            context="microsoft", message=mfa_message,
        )

    # Dismiss "Stay signed in?" (KMSI) if it appears. Some tenants skip it — after
    # MFA approval we may already be redirected back to the app.
    _dismiss_stay_signed_in(driver, wait_short, wait_long)

    # Login completed — let the MFA handler close any prompt it was showing.
    _resolve_mfa(mfa_handler)

    # Switch back to original window
    driver.switch_to.window(original_window)


def duo_login(driver: WebDriver, mfa_handler=None):
    # TODO: This is not working when in streamlit cloud. Need to get values set before this line
    # from cqc_cpcc.utilities.env_constants import INSTRUCTOR_USERID, INSTRUCTOR_PASS

    instructor_user_id = os.environ["INSTRUCTOR_USERID"]
    instructor_password = os.environ["INSTRUCTOR_PASS"]

    wait = get_driver_wait(driver, 3)  # Using shorter wait times

    original_window = driver.current_window_handle

    # Wait for title to change
    wait.until(EC.title_is("Web Login Service"))

    # Wait for login elements
    wait.until(
        lambda d: d.find_element(By.XPATH, "//div[@class='sr-only' and contains(text(),'Login')]"),
        "Waiting for login screen presence")

    # Login
    username_field = driver.find_element(By.ID, "username")
    password_field = driver.find_element(By.ID, "password")
    username_field.send_keys(instructor_user_id)
    password_field.send_keys(instructor_password)
    # login_field = driver.find_element(By.NAME, "_eventId_proceed")
    # login_field.click()
    click_element_wait_retry(driver, wait, "_eventId_proceed", "Waiting for login field", By.NAME)

    # Duo sends an automatic push to the instructor's device at this point. When a
    # number-matching prompt is shown, surface the number + screenshot through the
    # MFA handler (web app); otherwise fall back to opening the screenshot locally.
    _notify_mfa(
        driver,
        "duo",
        mfa_handler,
        "Duo push sent — approve it on your device "
        "(enter the number shown if prompted) to continue.",
    )

    # Switch to Duo Iframe
    # duo_frame = wait.until(lambda d: d.find_element(By.ID, "duo_iframe"), "Waiting for Duo Iframe")
    # wait.until(EC.frame_to_be_available_and_switch_to_it(duo_frame))

    # NOTE: Duo push happens automatically now. Used to require a button push
    # click_element_wait_retry(driver, wait, "//button[contains(text(),'Send Me a Push')]", "Waiting for auth buttons")

    try:
        # Click the no to is this your device message
        login_message = click_element_wait_retry(driver, wait,
                                                 "//button[contains(text(),'No, other people use this device')]",
                                                 "Waiting to click 'No, other people use this device' button",
                                                 max_try=1)

        # Wait until login accepted
        wait.until(
            EC.invisibility_of_element(login_message),
            'Waiting for login to be successful')
        _resolve_mfa(mfa_handler)
    except TimeoutException:
        # The "No, other people use this device" button did not appear in time.
        # This may mean Duo approval is still pending or an unexpected prompt
        # appeared.
        if mfa_handler is not None:
            # Web app: never block on stdin. Surface the prompt/screenshot to the
            # UI and let the caller wait for login to complete or be cancelled.
            _notify_mfa(
                driver,
                "duo",
                mfa_handler,
                "Duo authentication is still pending — approve the push "
                "(enter the number shown if prompted) on your device.",
            )
        elif not IS_GITHUB_ACTION:
            # CLI: take a screenshot and pause for manual intervention.
            take_and_show_screenshot(driver, "duo_timeout")
            wait_for_user_action(
                driver,
                "Duo authentication did not complete automatically.\n"
                "Please approve the Duo push on your device (or resolve any "
                "prompt shown in the screenshot above), then press Enter to continue.",
                take_screenshot=False,  # already taken above
            )
        _resolve_mfa(mfa_handler)

    # Switch back to original window
    driver.switch_to.window(original_window)
